from docx import Document

# Create a new document
document = Document()

# Add title
document.add_heading(
    "SMART QUESTION PAPER GENERATOR",
    level=1
)

# Add paragraphs
document.add_paragraph(
    "Department of Bachelor in Computer Applications"
)

document.add_paragraph(
    "Subject : Python Programming"
)

document.add_paragraph(
    "Maximum Marks : 20"
)

# Save document
document.save("Question_Paper.docx")

print("Document created successfully!")