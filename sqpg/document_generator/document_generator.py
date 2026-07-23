from docx import Document


def generate_document(config, file_name):
    """
    Generate a question paper document.
    """

    document = Document()

    # College Name
    document.add_heading(
        config["college_name"],
        level=1
    )

    # Department
    document.add_paragraph(
        config["department"]
    )

    # Maximum Marks
    document.add_paragraph(
        f"Maximum Marks : {config['default_marks']}"
    )

    document.save(file_name)

    print("Document Generated Successfully.")